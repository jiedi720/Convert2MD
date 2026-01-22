#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word (DOCX) 到 Markdown 转换器
Author: Your Name
Date: 2026-01-22
Description: 将 Word 文档转换为 Markdown 格式
"""

import docx
from pathlib import Path
import os
from typing import Dict, Any
from ..logger import app_logger


def convert_docx_to_md(docx_path: str, output_dir: str = None) -> bool:
    """
    将 DOCX 文件转换为 Markdown 格式
    :param docx_path: DOCX 文件路径
    :param output_dir: 输出目录，默认为 DOCX 文件所在目录
    :return: 转换是否成功
    """
    try:
        # 确定输出目录
        if output_dir is None:
            output_dir = os.path.dirname(docx_path)

        # 生成输出文件名
        docx_filename = Path(docx_path).stem
        output_path = os.path.join(output_dir, f"{docx_filename}.md")

        # 读取 DOCX 文件
        document = docx.Document(docx_path)

        # 提取内容并转换为 Markdown 格式
        md_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # 根据段落样式确定 Markdown 格式
            style = paragraph.style.name.lower()

            if 'heading' in style or style.startswith('标题'):
                # 确定标题级别
                level = 1  # 默认为一级标题
                if 'heading 1' in style or '标题1' in style:
                    level = 1
                elif 'heading 2' in style or '标题2' in style:
                    level = 2
                elif 'heading 3' in style or '标题3' in style:
                    level = 3
                elif 'heading 4' in style or '标题4' in style:
                    level = 4
                elif 'heading 5' in style or '标题5' in style:
                    level = 5
                elif 'heading 6' in style or '标题6' in style:
                    level = 6

                md_lines.append('#' * level + ' ' + text)
            else:
                # 普通段落
                md_lines.append(text)

        # 处理表格
        for table in document.tables:
            if table.rows:
                # 添加表头
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                if any(header_cells):
                    md_lines.append('| ' + ' | '.join(header_cells) + ' |')
                    md_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')

                    # 添加表格内容
                    for row in table.rows[1:]:
                        row_cells = [cell.text.strip() for cell in row.cells]
                        md_lines.append('| ' + ' | '.join(row_cells) + ' |')

        # 写入 Markdown 文件
        with open(output_path, 'w', encoding='utf-8') as md_file:
            md_file.write('\n\n'.join(md_lines))

        app_logger.info(f"DOCX 转换成功: {docx_path} -> {output_path}")
        return True

    except Exception as e:
        app_logger.error(f"DOCX 转换失败 {docx_path}: {e}")
        return False


def extract_text_and_structure(docx_path: str) -> dict:
    """
    提取 DOCX 的文本和结构信息
    :param docx_path: DOCX 文件路径
    :return: 包含文本和结构信息的字典
    """
    try:
        document = docx.Document(docx_path)

        content = {
            'title': '',
            'paragraphs': [],
            'lists': [],
            'tables': [],
            'metadata': {}
        }

        # 提取文档属性
        core_props = document.core_properties
        content['metadata'] = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else ''
        }

        # 如果没有标题，尝试从文档中获取
        if not content['title']:
            content['title'] = content['metadata']['title']

        # 提取段落和结构
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style = paragraph.style.name.lower()

            if 'heading' in style or style.startswith('标题'):
                content['paragraphs'].append({'text': text, 'type': 'heading', 'style': style})
            elif text.startswith(('*', '-', '+')) or text.startswith(tuple(str(i) + '.' for i in range(1, 10))):
                content['lists'].append(text)
            else:
                content['paragraphs'].append({'text': text, 'type': 'paragraph', 'style': style})

        # 提取表格
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                content['tables'].append(table_data)

        return content

    except Exception as e:
        app_logger.error(f"提取 DOCX 结构失败 {docx_path}: {e}")
        return {}


class DocxConverter:
    """
    DOCX 转换器实现类
    """

    def __init__(self, input_path: str, output_path: str = None):
        """
        初始化转换器
        :param input_path: 输入文件路径
        :param output_path: 输出文件路径，如果为None则自动生成
        """
        self.input_path = Path(input_path)
        self.output_path = Path(output_path) if output_path else self._generate_output_path()

        # 验证输入文件
        if not self.input_path.exists():
            raise FileNotFoundError(f"输入文件不存在: {input_path}")

        if not self.input_path.is_file():
            raise ValueError(f"输入路径不是文件: {input_path}")

    def _generate_output_path(self):
        """
        生成输出文件路径
        :return: 输出文件路径
        """
        input_stem = self.input_path.stem
        output_dir = self.input_path.parent
        return output_dir / f"{input_stem}.md"

    def convert(self) -> bool:
        """
        执行转换操作
        :return: 转换是否成功
        """
        try:
            # 确保输出目录存在
            self.output_path.parent.mkdir(parents=True, exist_ok=True)

            # 读取 DOCX 文件
            document = docx.Document(self.input_path)

            # 提取内容并转换为 Markdown 格式
            md_lines = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                # 根据段落样式确定 Markdown 格式
                style = paragraph.style.name.lower()

                if 'heading' in style or style.startswith('标题'):
                    # 确定标题级别
                    level = 1  # 默认为一级标题
                    if 'heading 1' in style or '标题1' in style:
                        level = 1
                    elif 'heading 2' in style or '标题2' in style:
                        level = 2
                    elif 'heading 3' in style or '标题3' in style:
                        level = 3
                    elif 'heading 4' in style or '标题4' in style:
                        level = 4
                    elif 'heading 5' in style or '标题5' in style:
                        level = 5
                    elif 'heading 6' in style or '标题6' in style:
                        level = 6

                    md_lines.append('#' * level + ' ' + text)
                else:
                    # 普通段落
                    md_lines.append(text)

            # 处理表格
            for table in document.tables:
                if table.rows:
                    # 添加表头
                    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                    if any(header_cells):
                        md_lines.append('| ' + ' | '.join(header_cells) + ' |')
                        md_lines.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')

                        # 添加表格内容
                        for row in table.rows[1:]:
                            row_cells = [cell.text.strip() for cell in row.cells]
                            md_lines.append('| ' + ' | '.join(row_cells) + ' |')

# 写入 Markdown 文件
            with open(self.output_path, 'w', encoding='utf-8') as md_file:
                md_file.write('\n\n'.join(md_lines))

            app_logger.info(f"DOCX 转换成功: {self.input_path} -> {self.output_path}")
            return True

        except Exception as e:
            app_logger.error(f"DOCX 转换失败 {self.input_path}: {e}")
            return False

    def extract_metadata(self) -> Dict[str, Any]:
        """
        提取文档元数据
        :return: 包含元数据的字典
        """
        try:
            document = docx.Document(self.input_path)

            # 提取文档属性
            core_props = document.core_properties
            return {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'category': core_props.category or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or ''
            }

        except Exception as e:
            app_logger.error(f"提取 DOCX 元数据失败 {self.input_path}: {e}")
            return {}

    def extract_structure(self) -> Dict[str, Any]:
        """
        提取文档结构信息
        :return: 包含结构信息的字典
        """
        try:
            document = docx.Document(self.input_path)

            content = {
                'title': '',
                'paragraphs': [],
                'lists': [],
                'tables': [],
                'images': [],
                'metadata': {},
                'styles': []
            }

            # 提取文档属性
            core_props = document.core_properties
            content['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else ''
            }

            # 如果没有标题，尝试从文档中获取
            if not content['title']:
                content['title'] = content['metadata']['title']

            # 提取段落和结构
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue

                style = paragraph.style.name.lower()

                if 'heading' in style or style.startswith('标题'):
                    content['paragraphs'].append({'text': text, 'type': 'heading', 'style': style})
                elif text.startswith(('*', '-', '+')) or text.startswith(tuple(str(i) + '.' for i in range(1, 10))):
                    content['lists'].append(text)
                else:
                    content['paragraphs'].append({'text': text, 'type': 'paragraph', 'style': style})

            # 提取表格
            for table in document.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    content['tables'].append(table_data)

            # 提取样式信息
            for style in document.styles:
                content['styles'].append({
                    'name': style.name,
                    'type': str(style.type),
                    'builtin': style.is_builtin
                })

            return content

        except Exception as e:
            app_logger.error(f"提取 DOCX 结构失败 {self.input_path}: {e}")
            return {}


if __name__ == "__main__":
    # 测试代码
    app_logger.info("DOCX 转换器测试")